# ============================================================================
# File: backend/ai-orchestration-layer/src/capabilities/document_parser.py
# DOCUMENT PARSING UTILITIES FOR RAG
# ============================================================================
# Supports: PDF, DOCX, TXT, MD, CSV files
# Includes text chunking with configurable overlap
# UPDATED: Added XML/HTML tag stripping and minimum chunk size filter
# ============================================================================

import os
import io
import re
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import logging

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    CSV = "csv"
    UNKNOWN = "unknown"


@dataclass
class ParsedChunk:
    """A chunk of parsed document content"""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Result of parsing a document"""
    doc_id: str
    filename: str
    doc_type: DocumentType
    raw_text: str
    chunks: List[ParsedChunk]
    metadata: Dict[str, Any]
    page_count: int = 1
    char_count: int = 0
    word_count: int = 0
    parsed_at: str = field(default_factory=lambda: datetime.utcnow().isoformat() + "Z")


class TextCleaner:
    """
    Cleans and preprocesses text before chunking.
    Handles XML/HTML tags, code artifacts, and normalizes whitespace.
    """

    # Common XML/HTML tag patterns
    TAG_PATTERN = re.compile(r'<[^>]+>', re.DOTALL)

    # CDATA sections
    CDATA_PATTERN = re.compile(r'<!\[CDATA\[.*?\]\]>', re.DOTALL)

    # XML declarations and processing instructions
    XML_DECL_PATTERN = re.compile(r'<\?[^>]+\?>', re.DOTALL)

    # HTML entities
    HTML_ENTITIES = {
        '&nbsp;': ' ',
        '&lt;': '<',
        '&gt;': '>',
        '&amp;': '&',
        '&quot;': '"',
        '&apos;': "'",
        '&#39;': "'",
        '&#x27;': "'",
        '&ldquo;': '"',
        '&rdquo;': '"',
        '&lsquo;': "'",
        '&rsquo;': "'",
        '&mdash;': '—',
        '&ndash;': '–',
        '&hellip;': '...',
    }

    # Numeric HTML entities pattern
    NUMERIC_ENTITY_PATTERN = re.compile(r'&#(\d+);')
    HEX_ENTITY_PATTERN = re.compile(r'&#x([0-9a-fA-F]+);')

    @classmethod
    def clean(cls, text: str, strip_tags: bool = True) -> str:
        """
        Clean text by removing XML/HTML tags and normalizing whitespace.

        Args:
            text: Raw text to clean
            strip_tags: Whether to strip XML/HTML tags (default True)

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        if strip_tags:
            # Remove CDATA sections (extract content)
            text = cls.CDATA_PATTERN.sub('', text)

            # Remove XML declarations
            text = cls.XML_DECL_PATTERN.sub('', text)

            # Remove all XML/HTML tags
            text = cls.TAG_PATTERN.sub(' ', text)

            # Decode HTML entities
            for entity, char in cls.HTML_ENTITIES.items():
                text = text.replace(entity, char)

            # Decode numeric entities
            def decode_numeric(match):
                try:
                    return chr(int(match.group(1)))
                except (ValueError, OverflowError):
                    return ''

            def decode_hex(match):
                try:
                    return chr(int(match.group(1), 16))
                except (ValueError, OverflowError):
                    return ''

            text = cls.NUMERIC_ENTITY_PATTERN.sub(decode_numeric, text)
            text = cls.HEX_ENTITY_PATTERN.sub(decode_hex, text)

        # Normalize whitespace
        # Replace multiple spaces/tabs with single space
        text = re.sub(r'[ \t]+', ' ', text)

        # Replace 3+ newlines with 2 newlines
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove spaces at the beginning/end of lines
        text = re.sub(r' *\n *', '\n', text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    @classmethod
    def looks_like_markup(cls, text: str) -> bool:
        """
        Detect if text appears to contain significant XML/HTML markup.

        Returns True if text has many tags, suggesting it should be cleaned.
        """
        if not text:
            return False

        tag_count = len(cls.TAG_PATTERN.findall(text))
        text_length = len(text)

        # If more than 1 tag per 100 chars, it's probably markup-heavy
        if text_length > 0 and tag_count > text_length / 100:
            return True

        # Check for common XML/HTML patterns
        markup_indicators = [
            '<?xml', '<!DOCTYPE', '<html', '<body', '<div', '<span',
            '<content>', '</content>', '<file>', '</file>',
            '<repo-to-text>', '</repo-to-text>',  # Common in code exports
        ]

        text_lower = text.lower()
        for indicator in markup_indicators:
            if indicator in text_lower:
                return True

        return False


class ChunkingStrategy:
    """
    Text chunking with configurable size and overlap.
    Uses recursive character splitting for better semantic boundaries.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 50,
        separators: Optional[List[str]] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size  # NEW: Minimum chunk size
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    def split_text(self, text: str) -> List[Tuple[str, int, int]]:
        """
        Split text into chunks with overlap.

        Returns:
            List of tuples: (chunk_text, start_char, end_char)
        """
        if not text:
            return []

        if len(text) <= self.chunk_size:
            stripped = text.strip()
            if len(stripped) >= self.min_chunk_size:
                return [(stripped, 0, len(text))]
            return []

        chunks = []
        start = 0

        while start < len(text):
            end = min(start + self.chunk_size, len(text))

            if end < len(text):
                # Find a good break point
                chunk_text = text[start:end]
                break_point = self._find_break_point(chunk_text)
                if break_point > 0:
                    end = start + break_point

            chunk_content = text[start:end].strip()

            # NEW: Only add chunk if it meets minimum size
            if len(chunk_content) >= self.min_chunk_size:
                chunks.append((chunk_content, start, end))

            # Move start with overlap
            start = max(start + 1, end - self.chunk_overlap)

        return chunks

    def _find_break_point(self, text: str) -> int:
        """Find the best break point in text based on separators."""
        min_position = int(self.chunk_size * 0.5)

        for separator in self.separators:
            if not separator:
                continue

            pos = text.rfind(separator)
            if pos >= min_position:
                return pos + len(separator)

        return len(text)


class DocumentParser:
    """
    Multi-format document parser.
    Extracts text and creates chunks for RAG indexing.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 50,
        auto_clean_markup: bool = True
    ):
        self.chunker = ChunkingStrategy(chunk_size, chunk_overlap, min_chunk_size)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.auto_clean_markup = auto_clean_markup
        self.cleaner = TextCleaner()

    def detect_type(self, filename: str) -> DocumentType:
        """Detect document type from filename."""
        if not filename:
            return DocumentType.UNKNOWN

        ext = os.path.splitext(filename.lower())[1].lstrip(".")

        type_map = {
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "doc": DocumentType.DOCX,
            "txt": DocumentType.TXT,
            "text": DocumentType.TXT,
            "md": DocumentType.MD,
            "markdown": DocumentType.MD,
            "csv": DocumentType.CSV,
        }

        return type_map.get(ext, DocumentType.UNKNOWN)

    def generate_doc_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID from filename and content hash."""
        content_hash = hashlib.md5(content).hexdigest()[:8]
        safe_name = "".join(c if c.isalnum() else "_" for c in filename)[:20]
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        return f"doc_{safe_name}_{content_hash}_{timestamp}"

    async def parse(
        self,
        filename: str,
        content: bytes,
        user_id: Optional[int] = None,
        extra_metadata: Optional[Dict[str, Any]] = None
    ) -> ParsedDocument:
        """
        Parse a document and extract text chunks.

        Args:
            filename: Original filename
            content: File content as bytes
            user_id: Optional user ID for metadata
            extra_metadata: Additional metadata to attach

        Returns:
            ParsedDocument with chunks
        """
        doc_type = self.detect_type(filename)
        doc_id = self.generate_doc_id(filename, content)

        # Extract text based on type
        if doc_type == DocumentType.PDF:
            raw_text, page_count = await self._parse_pdf(content)
        elif doc_type == DocumentType.DOCX:
            raw_text, page_count = await self._parse_docx(content)
        elif doc_type in (DocumentType.TXT, DocumentType.MD):
            raw_text, page_count = await self._parse_text(content)
        elif doc_type == DocumentType.CSV:
            raw_text, page_count = await self._parse_csv(content)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}. Supported: PDF, DOCX, TXT, MD, CSV")

        # NEW: Clean text if it contains markup
        cleaned_markup = False
        if self.auto_clean_markup and TextCleaner.looks_like_markup(raw_text):
            logger.info(f"Detected markup in {filename}, cleaning XML/HTML tags")
            raw_text = TextCleaner.clean(raw_text, strip_tags=True)
            cleaned_markup = True

        # Create chunks
        chunk_data = self.chunker.split_text(raw_text)
        chunks = []

        for i, (chunk_text, start, end) in enumerate(chunk_data):
            chunk_metadata = {
                "doc_id": doc_id,
                "filename": filename,
                "chunk_index": i,
                "total_chunks": len(chunk_data),
                "doc_type": doc_type.value,
            }
            if user_id is not None:
                chunk_metadata["user_id"] = user_id
            if extra_metadata:
                chunk_metadata.update(extra_metadata)

            chunks.append(ParsedChunk(
                content=chunk_text,
                chunk_index=i,
                start_char=start,
                end_char=end,
                metadata=chunk_metadata
            ))

        # Build document metadata
        metadata = {
            "filename": filename,
            "doc_type": doc_type.value,
            "page_count": page_count,
            "chunk_count": len(chunks),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "markup_cleaned": cleaned_markup,  # NEW: Track if we cleaned markup
        }
        if user_id is not None:
            metadata["user_id"] = user_id
        if extra_metadata:
            metadata.update(extra_metadata)

        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            doc_type=doc_type,
            raw_text=raw_text,
            chunks=chunks,
            metadata=metadata,
            page_count=page_count,
            char_count=len(raw_text),
            word_count=len(raw_text.split())
        )

    async def _parse_pdf(self, content: bytes) -> Tuple[str, int]:
        """Parse PDF content using pdfplumber or pypdf."""
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber

            text_parts = []
            page_count = 0

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())

            return "\n\n".join(text_parts), page_count

        except ImportError:
            pass

        # Fallback to pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text.strip())

            return "\n\n".join(text_parts), page_count

        except ImportError:
            raise ImportError("PDF parsing requires pdfplumber or pypdf. Install with: pip install pdfplumber pypdf")

    async def _parse_docx(self, content: bytes) -> Tuple[str, int]:
        """Parse DOCX content using python-docx."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            # Estimate page count (~3000 chars per page)
            page_count = max(1, len(full_text) // 3000)

            return full_text, page_count

        except ImportError:
            raise ImportError("DOCX parsing requires python-docx. Install with: pip install python-docx")

    async def _parse_text(self, content: bytes) -> Tuple[str, int]:
        """Parse plain text or markdown content."""
        # Try different encodings
        text = None
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            text = content.decode("utf-8", errors="replace")

        # Clean up text
        text = text.strip()

        # Estimate page count
        page_count = max(1, len(text) // 3000)
        return text, page_count

    async def _parse_csv(self, content: bytes) -> Tuple[str, int]:
        """Parse CSV content into readable text."""
        import csv

        text = content.decode("utf-8", errors="replace")

        try:
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
        except csv.Error:
            # If CSV parsing fails, treat as plain text
            return text, 1

        if not rows:
            return "", 1

        # Convert to readable format
        text_parts = []
        headers = rows[0] if rows else []

        for i, row in enumerate(rows[1:], 1):
            if headers and len(row) >= len(headers):
                row_text = "; ".join(
                    f"{h}: {v}"
                    for h, v in zip(headers, row)
                    if v.strip()
                )
            else:
                row_text = "; ".join(v for v in row if v.strip())

            if row_text:
                text_parts.append(f"Row {i}: {row_text}")

        full_text = "\n".join(text_parts)
        page_count = max(1, len(rows) // 50)

        return full_text, page_count


# Singleton instance
_parser_instance: Optional[DocumentParser] = None


def get_document_parser(
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    min_chunk_size: int = 50
) -> DocumentParser:
    """Get or create document parser instance."""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            min_chunk_size=min_chunk_size
        )
    return _parser_instance